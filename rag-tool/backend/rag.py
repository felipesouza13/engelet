"""
Núcleo do RAG: descoberta de arquivos, extração de texto, chunking,
embeddings (sentence-transformers, local) e persistência no Chroma.

Mantido independente do FastAPI para facilitar testes e reuso.
"""
from __future__ import annotations

import json
import os
import threading
from pathlib import Path
from typing import Callable, Iterable

# ---------------------------------------------------------------------------
# Caminhos base
# ---------------------------------------------------------------------------
RAIZ = Path(__file__).resolve().parents[2]          # .../engelet
LIVROS = RAIZ / "livros_normas"                     # fontes
BACKEND = Path(__file__).resolve().parent
ESTADO_PATH = BACKEND / "estado.json"
CHROMA_PATH = BACKEND / "chroma_db"

EXTENSOES = {".pdf", ".md", ".txt", ".docx"}

# Subpastas ignoradas na varredura (backups de originais escaneados pré-OCR).
PASTAS_IGNORADAS = {"originais-escaneados"}

# Modelo multilíngue (bom para PT-BR). e5 usa prefixos "passage:"/"query:".
MODELO_EMB = os.environ.get("RAG_EMB_MODEL", "intfloat/multilingual-e5-small")
CHUNK_SIZE = int(os.environ.get("RAG_CHUNK_SIZE", "1000"))      # caracteres
CHUNK_OVERLAP = int(os.environ.get("RAG_CHUNK_OVERLAP", "200"))  # caracteres
COLLECTION = "livros_normas"

# ---------------------------------------------------------------------------
# Estado (status de aprovação/indexação por arquivo) — persistido em JSON
# ---------------------------------------------------------------------------
# status possíveis: "pendente" | "aprovado" | "indexado"
_estado_lock = threading.Lock()


def _carregar_estado() -> dict:
    if ESTADO_PATH.exists():
        try:
            return json.loads(ESTADO_PATH.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
    return {}


def _salvar_estado(estado: dict) -> None:
    # gravação atômica: escreve em tmp e renomeia, evitando corromper o
    # estado.json se o processo for interrompido no meio da escrita.
    tmp = ESTADO_PATH.with_suffix(".json.tmp")
    tmp.write_text(
        json.dumps(estado, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    os.replace(tmp, ESTADO_PATH)


def _checar_extraivel(caminho: Path) -> dict:
    """
    Checagem leve: o arquivo tem camada de texto extraível para o RAG?

    Retorna {"extraivel": bool, "motivo": str, "chars": int}.
    PDFs escaneados (só imagem) retornam extraivel=False — precisariam de OCR.
    """
    ext = caminho.suffix.lower()
    try:
        if ext == ".pdf":
            import fitz  # PyMuPDF

            total = 0
            with fitz.open(caminho) as doc:
                n_paginas = doc.page_count
                # amostra as primeiras páginas (suficiente p/ detectar texto)
                for i in range(min(n_paginas, 5)):
                    total += len(doc[i].get_text().strip())
                    if total >= 200:
                        break
            if total >= 200:
                return {"extraivel": True, "motivo": "Texto detectado", "chars": total}
            return {
                "extraivel": False,
                "motivo": "PDF sem texto (provável escaneado — precisa OCR)",
                "chars": total,
            }
        if ext == ".docx":
            import docx

            doc = docx.Document(str(caminho))
            total = sum(len(p.text.strip()) for p in doc.paragraphs)
            if total > 0:
                return {"extraivel": True, "motivo": "Texto detectado", "chars": total}
            return {"extraivel": False, "motivo": "DOCX sem texto", "chars": 0}
        if ext in {".md", ".txt"}:
            amostra = caminho.read_text(encoding="utf-8", errors="ignore")[:2000].strip()
            if amostra:
                return {"extraivel": True, "motivo": "Texto detectado", "chars": len(amostra)}
            return {"extraivel": False, "motivo": "Arquivo vazio", "chars": 0}
    except Exception as e:  # noqa: BLE001
        return {"extraivel": False, "motivo": f"Erro ao ler: {e}", "chars": 0}
    return {"extraivel": False, "motivo": "Formato não suportado", "chars": 0}


def listar_arquivos() -> list[dict]:
    """Varre livros_normas recursivamente e mescla com o estado persistido."""
    LIVROS.mkdir(exist_ok=True)
    with _estado_lock:
        estado = _carregar_estado()
        encontrados: dict[str, dict] = {}
        for caminho in sorted(LIVROS.rglob("*")):
            if caminho.is_file() and caminho.suffix.lower() in EXTENSOES:
                rel = caminho.relative_to(LIVROS).as_posix()
                if rel.split("/", 1)[0] in PASTAS_IGNORADAS:
                    continue
                info = estado.get(rel, {})
                stat = caminho.stat()
                mtime = int(stat.st_mtime)

                # extraibilidade: usa cache se mtime não mudou; senão recalcula
                cache = info.get("extr")
                if not cache or cache.get("mtime") != mtime:
                    checagem = _checar_extraivel(caminho)
                    cache = {"mtime": mtime, **checagem}
                    info = {**info, "extr": cache}
                    estado[rel] = info

                encontrados[rel] = {
                    "path": rel,
                    "nome": caminho.name,
                    "ext": caminho.suffix.lower().lstrip("."),
                    "tamanho": stat.st_size,
                    "status": info.get("status", "pendente"),
                    "chunks": info.get("chunks", 0),
                    "extraivel": cache.get("extraivel", False),
                    "motivo_extr": cache.get("motivo", ""),
                }
        # mantém apenas arquivos que ainda existem, preservando campos
        estado_limpo: dict[str, dict] = {}
        for rel in encontrados:
            prev = estado.get(rel, {})
            estado_limpo[rel] = {
                "status": prev.get("status", "pendente"),
                "chunks": prev.get("chunks", 0),
                "extr": prev.get("extr", {}),
            }
        _salvar_estado(estado_limpo)
        return list(encontrados.values())


def definir_status(rel: str, status: str) -> None:
    with _estado_lock:
        estado = _carregar_estado()
        item = estado.setdefault(rel, {"status": "pendente", "chunks": 0})
        item["status"] = status
        _salvar_estado(estado)


def _registrar_indexado(rel: str, n_chunks: int) -> None:
    with _estado_lock:
        estado = _carregar_estado()
        item = estado.setdefault(rel, {"status": "pendente", "chunks": 0})
        item["status"] = "indexado"
        item["chunks"] = n_chunks
        _salvar_estado(estado)


def arquivos_aprovados() -> list[str]:
    return [a["path"] for a in listar_arquivos() if a["status"] == "aprovado"]


# ---------------------------------------------------------------------------
# Extração de texto por formato
# ---------------------------------------------------------------------------
def extrair_texto(caminho: Path) -> str:
    ext = caminho.suffix.lower()
    if ext == ".pdf":
        import fitz  # PyMuPDF

        partes = []
        with fitz.open(caminho) as doc:
            for pagina in doc:
                partes.append(pagina.get_text())
        return "\n".join(partes)
    if ext == ".docx":
        import docx

        documento = docx.Document(str(caminho))
        return "\n".join(p.text for p in documento.paragraphs)
    if ext in {".md", ".txt"}:
        return caminho.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Extensão não suportada: {ext}")


def chunk_texto(texto: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Chunking simples por caracteres com sobreposição, respeitando parágrafos."""
    texto = texto.replace("\r\n", "\n").strip()
    if not texto:
        return []
    chunks: list[str] = []
    inicio = 0
    n = len(texto)
    while inicio < n:
        fim = min(inicio + size, n)
        # tenta cortar em quebra de parágrafo/linha próxima ao fim
        if fim < n:
            corte = texto.rfind("\n", inicio + size - overlap, fim)
            if corte == -1:
                corte = texto.rfind(" ", inicio + size - overlap, fim)
            if corte != -1 and corte > inicio:
                fim = corte
        pedaco = texto[inicio:fim].strip()
        if pedaco:
            chunks.append(pedaco)
        if fim >= n:
            break
        inicio = max(fim - overlap, inicio + 1)
    return chunks


# ---------------------------------------------------------------------------
# Embeddings (lazy load do modelo) + Chroma
# ---------------------------------------------------------------------------
_model = None
_model_lock = threading.Lock()


def _get_model():
    global _model
    with _model_lock:
        if _model is None:
            from sentence_transformers import SentenceTransformer

            _model = SentenceTransformer(MODELO_EMB)
    return _model


def _get_collection():
    import chromadb

    client = chromadb.PersistentClient(path=str(CHROMA_PATH))
    return client.get_or_create_collection(
        name=COLLECTION, metadata={"hnsw:space": "cosine"}
    )


# ---------------------------------------------------------------------------
# Extração com progresso por página (para PDFs grandes dar feedback)
# ---------------------------------------------------------------------------
def _extrair_com_progresso(caminho: Path, emit, indice: int, total_arq: int, rel: str) -> str:
    ext = caminho.suffix.lower()
    if ext == ".pdf":
        import fitz  # PyMuPDF

        partes = []
        with fitz.open(caminho) as doc:
            n_pag = doc.page_count
            for p in range(n_pag):
                partes.append(doc[p].get_text())
                if p % 10 == 0 or p == n_pag - 1:
                    emit(
                        {
                            "tipo": "extraindo",
                            "path": rel,
                            "indice": indice,
                            "total_arquivos": total_arq,
                            "pagina": p + 1,
                            "paginas": n_pag,
                        }
                    )
        return "\n".join(partes)
    # demais formatos são rápidos
    return extrair_texto(caminho)


# ---------------------------------------------------------------------------
# Pipeline de indexação INCREMENTAL (arquivo por arquivo) com progresso
# ---------------------------------------------------------------------------
def indexar(
    rels: Iterable[str],
    progresso: Callable[[dict], None] | None = None,
) -> dict:
    """
    Para cada arquivo: extrai (progresso por página) -> chunka -> embeddings ->
    grava no Chroma, emitindo eventos de progresso para a barra do front.

    Eventos:
      {"tipo": "inicio", "total_arquivos": n}
      {"tipo": "extraindo", "path", "indice", "total_arquivos", "pagina", "paginas"}
      {"tipo": "arquivo", "path", "indice", "total_arquivos", "chunks"}
      {"tipo": "chunk", "path", "chunk_no_arquivo", "chunks_no_arquivo",
                        "arquivos_feitos", "total_arquivos", "chunks_acumulados"}
      {"tipo": "arquivo_fim", "path", "arquivos_feitos", "total_arquivos", "chunks_acumulados"}
      {"tipo": "erro", "path", "msg"}
      {"tipo": "fim", "total_chunks", "arquivos"}
    """
    rels = list(rels)

    def emit(evt: dict):
        if progresso:
            progresso(evt)

    total_arq = len(rels)
    emit({"tipo": "inicio", "total_arquivos": total_arq})
    if total_arq == 0:
        emit({"tipo": "fim", "total_chunks": 0, "arquivos": 0})
        return {"total_chunks": 0, "arquivos": 0}

    model = _get_model()
    collection = _get_collection()

    batch = 32
    chunks_acum = 0
    arquivos_feitos = 0

    for i, rel in enumerate(rels):
        caminho = LIVROS / rel
        try:
            emit({"tipo": "extraindo", "path": rel, "indice": i + 1,
                  "total_arquivos": total_arq, "pagina": 0, "paginas": 0})
            texto = _extrair_com_progresso(caminho, emit, i + 1, total_arq, rel)
            pedacos = chunk_texto(texto)
        except Exception as e:  # noqa: BLE001
            emit({"tipo": "erro", "path": rel, "msg": str(e)})
            arquivos_feitos += 1
            emit({"tipo": "arquivo_fim", "path": rel, "arquivos_feitos": arquivos_feitos,
                  "total_arquivos": total_arq, "chunks_acumulados": chunks_acum})
            continue

        n = len(pedacos)
        emit({"tipo": "arquivo", "path": rel, "indice": i + 1,
              "total_arquivos": total_arq, "chunks": n})

        # remove versões anteriores deste arquivo (reindexação limpa)
        try:
            collection.delete(where={"path": rel})
        except Exception:  # noqa: BLE001
            pass

        feitos_arq = 0
        for ini in range(0, n, batch):
            lote = pedacos[ini : ini + batch]
            entradas = [f"passage: {t}" for t in lote]
            vetores = model.encode(
                entradas, normalize_embeddings=True, show_progress_bar=False
            ).tolist()
            ids = [f"{rel}::{ini + j}" for j in range(len(lote))]
            metadados = [{"path": rel, "chunk_idx": ini + j} for j in range(len(lote))]
            collection.add(
                ids=ids, embeddings=vetores, documents=lote, metadatas=metadados
            )
            feitos_arq += len(lote)
            emit({"tipo": "chunk", "path": rel,
                  "chunk_no_arquivo": feitos_arq, "chunks_no_arquivo": n,
                  "arquivos_feitos": arquivos_feitos, "total_arquivos": total_arq,
                  "chunks_acumulados": chunks_acum + feitos_arq})

        chunks_acum += n
        arquivos_feitos += 1
        _registrar_indexado(rel, n)
        emit({"tipo": "arquivo_fim", "path": rel, "arquivos_feitos": arquivos_feitos,
              "total_arquivos": total_arq, "chunks_acumulados": chunks_acum})

    emit({"tipo": "fim", "total_chunks": chunks_acum, "arquivos": arquivos_feitos})
    return {"total_chunks": chunks_acum, "arquivos": arquivos_feitos}


def consultar(pergunta: str, k: int = 5) -> list[dict]:
    """Busca semântica simples (útil para testar o índice)."""
    model = _get_model()
    collection = _get_collection()
    vetor = model.encode(
        [f"query: {pergunta}"], normalize_embeddings=True
    ).tolist()
    res = collection.query(query_embeddings=vetor, n_results=k)
    saidas = []
    for doc, meta, dist in zip(
        res["documents"][0], res["metadatas"][0], res["distances"][0]
    ):
        saidas.append({"texto": doc, "path": meta["path"], "distancia": dist})
    return saidas
